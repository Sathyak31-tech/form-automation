#!/usr/bin/env python3
"""
Smart Form Populator
Intelligently fills forms by understanding their structure and context.
Pure pattern-based approach without LLM dependencies.
"""

from __future__ import annotations
import json, os, re
from datetime import datetime
from typing import Dict, List, Tuple, Optional, Union
from docx.oxml import OxmlElement
from docx import Document
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.shared import Inches


def today_str(fmt: str = "%d-%m-%Y") -> str:
    return datetime.now().strftime(fmt)


def get_nested_value(d: dict, key: str, default: str = "") -> str:
    """Return a scalar or 'value' from a dict like {'value':..., 'iso':..., 'raw':...}."""
    if not isinstance(d, dict):
        return default
    v = d.get(key)
    if v is None:
        return default
    if isinstance(v, dict):
        return (v.get("raw") or v.get("value") or v.get("iso") or default) or default
    return str(v)


class SmartFormPopulator:
    # -------- Label regex (broad & forgiving) --------
    # -------- Label regex (broad & forgiving) --------
    NAME_LABEL_RE   = re.compile(
        r"\b("
        r"name(?!.*father)"                           # plain "Name" but not "Father Name"
        r"|employee\s*name"
        r"|member\s*name"
        r"|I\s*name"
        r"|applicant\s*name"
        r"|name\s+of\s+(?:the\s+)?employee(?:\s+in\s+full)?"
        r")\b",
        re.IGNORECASE,
    )

    FATHER_LABEL_RE = re.compile(
        r"\b("
        r"father[’']?s?\s*name"
        r"|husband[’']?s?\s*name"
        r"|guardian[’']?s?\s*name"
        r"|s\/o"                                      # S/o
        r"|d\/o"                                      # D/o
        r"|w\/o"                                      # W/o
        r"|son\s*\/\s*daughter\s*\/\s*wife\s*of"
        r")\b",
        re.IGNORECASE,
    )
    
    # Lines that start like "I, … s/o or d/o …" (declaration formats)
    DECLARATION_NAME_LINE_RE = re.compile(r"^\s*i\s*,", re.IGNORECASE)
    REL_MARKER_RE = re.compile(r"\b(s\/o|d\/o|w\/o|son\s*\/\s*daughter\s*\/\s*wife\s*of)\b", re.IGNORECASE)

    
    EMAIL_LABEL_RE  = re.compile(
        r"\b((personal|official)\s+)?e[-\s]?mail(\s*id)?\b|\bmail\s*id\b",
        re.IGNORECASE,
    )

    ADDRESS_LABEL_RE = re.compile(
        r"\b(?:(present|current|residential|communication|correspondence|registered|permanent|previous)\s+)?address\b",
        re.IGNORECASE,
    )

    DATE_LABEL_RE   = re.compile(r"\bdate\b", re.IGNORECASE)

    def __init__(self, data_file: str):
        with open(data_file, "r", encoding="utf-8") as f:
            data = json.load(f)

        self.form_fields: dict = data.get("form_fields", data)

        # quick references
        self.pd: dict = self.form_fields.get("personal_details", {}) or {}
        self.emp: list = self.form_fields.get("employment_history", []) or []
        self.edu: dict = self.form_fields.get("education", self.form_fields.get("education_history", {})) or {}
        self.addr: dict = self.form_fields.get("address_history", {}) or {}
        self.refs: list = self.form_fields.get("references", []) or []
        self.gaps: Union[dict, list] = self.form_fields.get("gaps", {}) or {}
        
        # Signature image path
        self.signature_image_path: Optional[str] = self.pd.get("signature_image_path") or os.environ.get("SIGNATURE_IMAGE_PATH")
        
        # Debug: Print signature path if found
        if self.signature_image_path:
            print(f"📷 Signature image path found: {self.signature_image_path}")
            if os.path.exists(self.signature_image_path):
                print(f"✅ Signature image file exists")
            else:
                print(f"⚠️  Warning: Signature image file does NOT exist at: {self.signature_image_path}")
        else:
            print(f"⚠️  No signature image path provided")

    # ----------------------------
    # Structure extraction helpers
    # ----------------------------
    
        
        # import re
    @staticmethod
    def _norm(s: str) -> str:
        return re.sub(r"\s+", " ", (s or "").strip().lower())

    @staticmethod
    def _is_placeholder(s: str) -> bool:
        if s is None: return True
        st = s.strip()
        return (st == "") or bool(re.fullmatch(r"[_\.\-\s–—]+", st))

    @staticmethod
    def _fill_after_colon(text: str, value: str):
        """Return (new_text, changed) with value filled after colon or placeholder."""
        if text is None:
            return text, False
        if ":" in text:
            left, right = text.split(":", 1)
            if SmartFormPopulator._is_placeholder(right):
                return left + ": " + str(value), True
        m = re.search(r"([:_：]\s*)[_\.\-\s–—]+$", text)
        if m:
            start = m.start()
            return text[:start] + m.group(1) + str(value), True
        return text, False

    @staticmethod
    def _tick(text: str, chosen: str, options=("Male","Female","Transgender")):
        if not text: return text, False
        v = (chosen or "").strip().lower()
        before = text
        for opt in options:
            text = re.sub(rf"\b{re.escape(opt)}\b",
                        lambda m: ("☑ " if m.group(0).strip().lower()==v else "☐ ") + m.group(0),
                        text, count=1, flags=re.IGNORECASE)
        return text, (text != before)
    
    @staticmethod
    def _has_keys(text: str, keys):
        t = SmartFormPopulator._norm(text)
        return any(SmartFormPopulator._norm(k) in t for k in keys)

    @staticmethod
    def _pxml_get_text(p_el) -> str:
        return "".join([(t.text or "") for t in p_el.iter() if getattr(t, "tag", "").endswith("}t")])

    @staticmethod
    def _pxml_set_text(p_el, new_text: str):
        # Replace entire paragraph XML with a single run + text
        for child in list(p_el):
            p_el.remove(child)
        r = OxmlElement("w:r")
        t = OxmlElement("w:t")
        t.text = new_text
        r.append(t)
        p_el.append(r)
    
    def _insert_signature_image(self, paragraph) -> bool:
        """Insert signature image into a paragraph if signature image path exists."""
        if not self.signature_image_path:
            return False
        
        if not os.path.exists(self.signature_image_path):
            print(f"Warning: Signature image path does not exist: {self.signature_image_path}")
            return False
        
        try:
            # Clear existing text/runs
            paragraph.clear()
            
            # Add the signature image
            run = paragraph.add_run()
            run.add_picture(self.signature_image_path, width=Inches(2.0))  # 2 inches wide
            print(f"✅ Inserted signature image from {self.signature_image_path}")
            return True
        except Exception as e:
            print(f"Warning: Could not insert signature image: {e}")
            import traceback
            print(traceback.format_exc())
            return False
    
    def _cell_has_image(self, cell) -> bool:
        """Check if a cell already contains an image."""
        try:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    # Check if run has an image (drawing element)
                    if run._element.xpath('.//a:blip') or run._element.xpath('.//pic:pic'):
                        return True
            return False
        except:
            return False
    
    def _insert_signature_in_cell(self, cell) -> bool:
        """Insert signature image into a table cell."""
        if not self.signature_image_path:
            return False
        
        if not os.path.exists(self.signature_image_path):
            print(f"Warning: Signature image path does not exist: {self.signature_image_path}")
            return False
        
        try:
            # Don't overwrite if cell already has an image
            if self._cell_has_image(cell):
                return False
            
            # Clear ALL content from the cell (text, runs, everything)
            # This is important because the cell might contain "Signature" as text
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    paragraph._element.remove(run._element)
            
            # Clear all paragraphs except the first one
            while len(cell.paragraphs) > 1:
                cell._element.remove(cell.paragraphs[-1]._element)
            
            # Get or create the first paragraph
            if len(cell.paragraphs) == 0:
                paragraph = cell.add_paragraph()
            else:
                paragraph = cell.paragraphs[0]
            
            # Clear any remaining content
            paragraph.clear()
            
            # Add signature image
            run = paragraph.add_run()
            run.add_picture(self.signature_image_path, width=Inches(2.0))  # 2 inches wide
            print(f"✅ Inserted signature image in cell from {self.signature_image_path}")
            return True
        except Exception as e:
            print(f"Warning: Could not insert signature image in cell: {e}")
            import traceback
            print(traceback.format_exc())
            return False
        
    DECLARATION_KEYWORDS = [
    (["name", "member name", "full name", "candidate name"], "name"),
    (["father", "father’s", "father's", "husband", "husband’s", "husband's"], "father_name"),
    (["date of birth", "dob", "d.o.b"], "date_of_birth"),
    (["email", "email id", "e-mail"], "email"),
    (["nationality"], "nationality"),
    (["pan", "permanent account number"], "pan_card"),
    (["aadhaar", "aadhar", "npr"], "aadhar_card"),
    (["passport"], "passport_details"),
    (["passport details"], "passport_details"),
]
    
    
    def _build_decl_values(self) -> dict:
        pd = self.form_fields.get("personal_details", {})
        dob = pd.get("date_of_birth", {})
        if isinstance(dob, dict):
            dob = dob.get("value", "")
        return {
            "name": pd.get("name",""),
            "father_name": pd.get("father_name",""),
            "email": pd.get("email",""),
            "date_of_birth": dob or pd.get("date_of_birth","") or "",
            "nationality": pd.get("nationality",""),
            "pan_card": pd.get("pan_card",""),
            "aadhar_card": pd.get("aadhar_card",""),
            "passport_no": pd.get("passport_no",""),
            "gender": pd.get("gender",""),
        }

    def _fill_declaration_form(self, doc) -> int:
    # """Fills EPFO-style Declaration form across paragraphs, tables, text boxes, and content controls."""
        values = self._build_decl_values()
        filled = 0
        filled_fields = set()  # Track what we've filled to avoid duplicates

        # A) Normal paragraphs
        for p in doc.paragraphs:
            txt = p.text or ""
            if any(w in txt.lower() for w in ["gender", "male", "female", "transgender"]):
                new_txt, ch = SmartFormPopulator._tick(txt, values.get("gender",""))
                if ch:
                    for r in p.runs: r.clear()
                    p.add_run(new_txt); filled += 1
                    continue
            for keys, vkey in SmartFormPopulator.DECLARATION_KEYWORDS:
                v = values.get(vkey, "")
                if not v: continue
                if SmartFormPopulator._has_keys(txt, keys):
                    new_txt, ch = SmartFormPopulator._fill_after_colon(txt, v)
                    if ch:
                        for r in p.runs: r.clear()
                        p.add_run(new_txt); filled += 1
                        filled_fields.add(vkey)
                    break
        
        # A2) Fill empty paragraphs near label paragraphs
        paragraphs = doc.paragraphs
        for i, p in enumerate(paragraphs):
            txt = p.text.strip()
            # Check if this is an empty paragraph that might be an input field
            if not txt or txt == "":
                # Look at nearby paragraphs (within 3 before) for labels
                for j in range(max(0, i-3), i):
                    label_text = paragraphs[j].text.strip()
                    # Try to match this label to a keyword
                    for keys, vkey in SmartFormPopulator.DECLARATION_KEYWORDS:
                        if vkey not in filled_fields and SmartFormPopulator._has_keys(label_text, keys):
                            v = values.get(vkey, "")
                            if v:
                                for r in p.runs: r.clear()
                                p.add_run(str(v))
                                filled += 1
                                filled_fields.add(vkey)
                                break
                    if p.text.strip():  # If we filled it, move to next empty paragraph
                        break

        # B) Tables
        for tbl in doc.tables:
            for row in tbl.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        txt = p.text or ""
                        if any(w in txt.lower() for w in ["gender", "male", "female", "transgender"]):
                            new_txt, ch = SmartFormPopulator._tick(txt, values.get("gender",""))
                            if ch:
                                for r in p.runs: r.clear()
                                p.add_run(new_txt); filled += 1
                                continue
                        for keys, vkey in SmartFormPopulator.DECLARATION_KEYWORDS:
                            v = values.get(vkey, "")
                            if not v: continue
                            if SmartFormPopulator._has_keys(txt, keys):
                                new_txt, ch = SmartFormPopulator._fill_after_colon(txt, v)
                                if ch:
                                    for r in p.runs: r.clear()
                                    p.add_run(new_txt); filled += 1
                                break

        # C) Text boxes & content controls (python-docx doesn't expose; traverse XML)
        body_el = doc.element.body
        p_elements = []
        for el in body_el.iter():
            tag = getattr(el, "tag", "")
            if tag.endswith("}txbxContent") or tag.endswith("}sdtContent"):
                for sub in el.iter():
                    if getattr(sub, "tag", "").endswith("}p"):
                        p_elements.append(sub)

        for p_el in p_elements:
            txt = SmartFormPopulator._pxml_get_text(p_el)
            if not txt:
                continue
            if any(w in txt.lower() for w in ["gender", "male", "female", "transgender"]):
                new_txt, ch = SmartFormPopulator._tick(txt, values.get("gender",""))
                if ch:
                    SmartFormPopulator._pxml_set_text(p_el, new_txt); filled += 1
                    continue
            for keys, vkey in SmartFormPopulator.DECLARATION_KEYWORDS:
                v = values.get(vkey, "")
                if not v: continue
                if SmartFormPopulator._has_keys(txt, keys):
                    new_txt, ch = SmartFormPopulator._fill_after_colon(txt, v)
                    if ch:
                        SmartFormPopulator._pxml_set_text(p_el, new_txt); filled += 1
                    break

        return filled
    
    


    def _row_text(self, row) -> str:
        """Return normalized text of the whole row (all cells joined)."""
        try:
            parts = []
            for c in row.cells:
                parts.append((c.text or "").strip())
            return self._norm(" | ".join(parts))
        except Exception:
            return ""

    def _row_label(self, row) -> str:
        """Whole-row text (normalized). Used by reference/gap filling and loose matching."""
        return self._row_text(row)

    def _table_has_any_text(self, table, keys, where: str = "left") -> bool:
        """
        True if ANY normalized key appears in the chosen part of the table.
        where='left' → check first cell of each row (typical label column)
        where='row'  → check the whole row (good for merged/3-col tables)
        """
        keys = [self._norm(k) for k in keys]

        for row in table.rows:
            if where == "left":
                text = ""
                if len(row.cells) >= 1:
                    text = self._norm(row.cells[0].text)
                else:
                    text = self._row_text(row)
            else:
                text = self._row_text(row)

            if any(k in text for k in keys):
                return True
        return False

    def _linearize_blocks(self, doc: Document):
        """Return a linear list of ('p', idx) or ('t', idx) in document order."""
        blocks = []
        p_idx = 0
        t_idx = 0
        for child in doc._element.body.iterchildren():
            if isinstance(child, CT_P):
                blocks.append(("p", p_idx))
                p_idx += 1
            elif isinstance(child, CT_Tbl):
                blocks.append(("t", t_idx))
                t_idx += 1
        return blocks

    def extract_form_structure(self, docx_path: str) -> Dict:
        doc = Document(docx_path)

        structure = {
            "paragraphs": [],
            "tables": [],
            "form_type": "unknown",
        }

        # collect paragraphs
        for i, para in enumerate(doc.paragraphs):
            text = (para.text or "").strip()
            if text:
                structure["paragraphs"].append({
                    "index": i,
                    "text": text,
                    "is_field": self.is_fillable_field(text),
                    "field_type": self.get_field_type(text),
                })

        # compute paragraph heading immediately preceding each table
        order = self._linearize_blocks(doc)
        last_para_for_table = {}
        last_seen_para = None
        for kind, idx in order:
            if kind == "p":
                last_seen_para = idx
            elif kind == "t":
                last_para_for_table[idx] = last_seen_para

        # collect tables with their heading
        for table_idx, table in enumerate(doc.tables):
            heading_text = ""
            near_para_idx = last_para_for_table.get(table_idx)
            if near_para_idx is not None and 0 <= near_para_idx < len(doc.paragraphs):
                heading_text = (doc.paragraphs[near_para_idx].text or "").strip()

            table_data = {
                "index": table_idx,
                "rows": [],
                "heading": heading_text,  # used to route Highest/Previous/Professional
            }
            for row_idx, row in enumerate(table.rows):
                row_data = {"index": row_idx, "cells": []}
                for cell_idx, cell in enumerate(row.cells):
                    text = (cell.text or "").strip()
                    row_data["cells"].append({
                        "index": cell_idx,
                        "text": text,
                        "is_field": self.is_fillable_field(text),
                        "field_type": self.get_field_type(text),
                    })
                table_data["rows"].append(row_data)

            structure["tables"].append(table_data)

        structure["form_type"] = self.determine_form_type(structure)
        return structure

    def is_fillable_field(self, text: str) -> bool:
        pats = [
            r":\s*$", r"\(Complete\)\*:", r"\(if any\)\*:",
            r"\bPrint Name\b", r"\bSignature\b", r"\bDate\b", r"\bTitle\b",
            r"Name.*:", r"Address.*:", r"Email.*:", r"Phone.*:",
        ]
        return any(re.search(p, text, re.I) for p in pats)

    def get_field_type(self, text: str) -> str:
        t = text.lower()
        if "background verification" in t:
            return "bgv_header"
        if "name" in t and "complete" in t and "father" not in t:
            return "full_name"
        if "father" in t and "name" in t:
            return "father_name"
        if "email" in t:
            return "email"
        if "address" in t:
            return "address"
        if "pan" in t:
            return "pan_card"
        if "aadhaar" in t or "aadhar" in t:
            return "aadhar_card"
        if "phone" in t or "mobile" in t:
            return "phone"
        if "date" in t and "birth" in t:
            return "date_of_birth"
        if "gender" in t or re.search(r"\bsex\b", t):
            return "gender"
        if "nationality" in t:
            return "nationality"
        if "signature" in t:
            return "signature"
        if "print name" in t:
            return "print_name"
        if re.fullmatch(r"\s*date\s*:?\s*", text, re.I):
            return "date"
        if "din" in t:
            return "din"
        if "passport" in t and "details" in t:
            return "passport_details"
        if "passport" in t and "issue" in t:
            return "passport_issue_date"
        if "passport" in t and "expiry" in t:
            return "passport_expiry_date"
        if "passport" in t:
            return "passport_no"
        # Handle passport date fields that don't explicitly mention "passport"
        if "issue date" in t and t.endswith(':'):
            return "passport_issue_date"
        if "expiry date" in t and "employment history" in t:
            return "passport_expiry_date_with_employment"
        if "expiry date" in t and t.endswith(':'):
            return "passport_expiry_date"
        return "unknown"

    def determine_form_type(self, structure: Dict) -> str:
        all_text = " ".join(p["text"] for p in structure["paragraphs"])
        for table in structure["tables"]:
            for row in table["rows"]:
                for cell in row["cells"]:
                    all_text += " " + cell["text"]
        t = all_text.lower()
        if "background verification" in t:
            return "background_verification"
        if "form 2 revised" in t or "nomination and declaration form" in t or "epf" in t and "nomination" in t:
            return "epf_nomination"
        if "gratuity" in t:
            return "gratuity"
        if "declaration for pf account linking with aadhar" in t:
            return "pf_account"
        if "declaration" in t:
            return "declaration"
        if "non-disclosure" in t or "nda" in t:
            return "nda"
        if "leave of absence" in t or re.search(r"\bloa\b", t):
            return "loa"
        return "general"

    # ----------------------------
    # Value mapping (robust)
    # ----------------------------
    def _get(self, *paths: str, default: str = "") -> str:
        """Try multiple keys in personal_details."""
        for k in paths:
            v = self.pd.get(k)
            if v:
                if isinstance(v, dict):
                    return v.get("raw") or v.get("value") or v.get("iso") or default
                return str(v)
        return default

    def _current_address_value(self) -> str:
        cur = (self.addr.get("current") or {})
        return cur.get("town_or_city_name", "") or ""

    def _permanent_address_value(self) -> str:
        perm = (self.addr.get("permanent") or {})
        return perm.get("town_or_city_name", "") or ""

    def _previous_address_value(self) -> str:
        prev = (self.addr.get("previous") or {})
        return prev.get("town_or_city_name", "") or ""

    def get_field_value(self, field_type: str) -> str:
        address_cur = self.addr.get("current", {}) or {}
        address_perm = self.addr.get("permanent", {}) or {}
        address_any = address_cur or address_perm
        mapping = {
            "full_name": self._get("name"),
            "father_name": self._get("father_name", "fathers_name"),
            "email": self._get("email"),
            "address": address_any.get("town_or_city_name", ""),
            "pan_card": self._get("pan_card", "pan_card_no"),
            "aadhar_card": self._get("aadhar_card", "aadhar_card_no"),
            "phone": address_any.get("phone_number", ""),
            "date_of_birth": self._get("date_of_birth"),
            "gender": self._get("gender"),
            "nationality": self._get("nationality"),
            "signature": "",  # keep blank for real signature
            "print_name": self._get("name"),
            "title": (self.emp[0].get("position_and_department", "") if self.emp else ""),
            "date": today_str(),
            "din": self._get("din"),
            "passport_no": self._get("passport_no"),
            "passport_issue_date": self._get("passport_issue_date"),
            "passport_expiry_date": self._get("passport_expiry_date"),
            "passport_expiry_date_with_employment": self._get("passport_expiry_date"),
            "passport_details": self._get_passport_details(),
        }
        return mapping.get(field_type, "")

    def _get_passport_details(self) -> str:
        """Get passport number only (dates are filled separately to avoid duplication)."""
        passport_no = self._get("passport_no")
        
        if not passport_no:
            return ""
        
        return f"Passport No: {passport_no}"

    # ----------------------------
    # Table classification helpers
    # ----------------------------
    def _looks_like_employment(self, table) -> bool:
        keys = ["employers name", "employer name", "employers name & branch", "employer name & branch", "position", "employment period",
                "employee code", "last salary", "reporting manager", "verify",
                "agency", "office", "telephone", "landline"]
        return self._table_has_any_text(table, keys, where="left")

    def _looks_like_education(self, table) -> bool:
        keys = ["university and college", "degree", "diploma", "course",
                "period of the course", "roll", "registration", "seat", "location"]
        return self._table_has_any_text(table, keys, where="left")

    def _looks_like_address(self, table) -> bool:
        keys = ["complete", "address", "town/ city", "town", "city",
                "duration of stay", "phone number", "current address",
                "previous address", "permanent address"]
        return (self._table_has_any_text(table, keys, where="left") or
                self._table_has_any_text(table, keys, where="row"))

    def _looks_like_reference(self, table) -> bool:
        # Look for specific reference section indicators
        keys = ["reference", "referee", "details of professional references"]
        if self._table_has_any_text(table, keys, where="row"):
            return True
        
        # Also look for reference-specific patterns like "Reference 1", "Reference 2"
        for row in table.rows:
            if len(row.cells) >= 1:
                text = row.cells[0].text.strip().lower()
                if re.search(r"\breference\s*(1|\(i\)|i|2|\(ii\)|ii)\b", text):
                    return True
        
        return False

    def _looks_like_gap(self, table) -> bool:
        keys = ["reason for gap", "period of gap", "address stayed during the gap",
                "gap", "career gap", "employment gap", "unemployment",
                "from", "to", "duration"]
        return self._table_has_any_text(table, keys, where="row")

    def _classify_table(self, table) -> str:
        if self._looks_like_employment(table):
            return "employment"
        if self._looks_like_education(table):
            return "education"
        
        # Check for mixed tables (address + reference + gap)
        has_address = self._looks_like_address(table)
        has_reference = self._looks_like_reference(table)
        has_gap = self._looks_like_gap(table)
        
        # If it has both address and reference/gap, classify as reference (mixed table)
        if has_reference and (has_address or has_gap):
            return "reference"
        elif has_address:
            return "address"
        elif has_reference:
            return "reference"
        elif has_gap:
            return "gap"
        
        return "unknown"

    def _blank(self, v) -> str:
        if v is None:
            return ""
        s = str(v).strip()
        return "" if s.upper() in {"", "N/A", "NA", "NONE", "NULL"} else s
    
    def _extract_city_name(self, full_address):
        """Extract city name from full address string."""
        if not full_address:
            return ""
        
        # Split by comma and take the last part (usually city, state)
        parts = full_address.split(',')
        if len(parts) >= 2:
            # Take the second-to-last part (city) and last part (state)
            city_state = parts[-2].strip() + ', ' + parts[-1].strip()
            return city_state
        elif len(parts) == 1:
            # If no comma, return the whole string
            return full_address.strip()
        else:
            return ""

    def _safe_join(self, parts, sep=" "):
        vals = [self._blank(p) for p in parts if self._blank(p)]
        return sep.join(vals)

    def _clear_right_cells(self, table):
        """Clear the LAST cell in each row (supports 2- or 3-col tables)."""
        for row in table.rows:
            if len(row.cells) >= 2:
                row.cells[-1].text = ""

    def _fill_cell_if_label(self, row, must_contain: List[str], value: str, fixes_ref: List[int]):
        if len(row.cells) < 2:
            return
        left = (row.cells[0].text or "").strip().lower()
        if all(k.lower() in left for k in must_contain):
            row.cells[-1].text = self._blank(value)
            fixes_ref[0] += 1

    # ----------------------------
    # Fillers
    # ----------------------------
    def _is_company_side(self, cell_text: str, row_text: str = "") -> bool:
        """Check if this is the company's side (ACCOLITE) vs user's side."""
        text_lower = (cell_text + " " + row_text).lower()
        # Company indicators
        company_keywords = ["accolite", "company", "employer", "organization"]
        # User indicators  
        user_keywords = ["you", "employee", "candidate"]
        
        # If it mentions company keywords and not user keywords, it's company side
        has_company = any(kw in text_lower for kw in company_keywords)
        has_user = any(kw in text_lower for kw in user_keywords)
        
        if has_company and not has_user:
            return True
        return False
    
    def _is_user_side(self, cell_text: str, row_text: str = "", name: str = "") -> bool:
        """Check if this is the user's side (You/Employee) vs company's side."""
        text_lower = (cell_text + " " + row_text).lower()
        name_lower = name.lower() if name else ""
        
        # User indicators
        user_keywords = ["you", "employee", "candidate"]
        # Company indicators
        company_keywords = ["accolite", "company", "employer", "organization"]
        
        # Check if name appears (user's side)
        if name and name_lower in text_lower:
            return True
        
        # If it mentions user keywords and not company keywords, it's user side
        has_user = any(kw in text_lower for kw in user_keywords)
        has_company = any(kw in text_lower for kw in company_keywords)
        
        if has_user and not has_company:
            return True
        
        # Default: if we can't determine, assume it's user side if it's the rightmost column
        return False

    def _fill_signatures_everywhere(self, doc: Document) -> int:
        """Fill signature fields in all paragraphs and tables."""
        fixes = 0
        name = self._get("name")
        
        # Scan all paragraphs for signature fields
        for p in doc.paragraphs:
            text = (p.text or "").strip()
            text_lower = text.lower()
            # Look for various signature patterns
            if (text_lower == "signature:" or 
                text_lower == "signature" or
                text_lower.endswith("signature:") or
                (":" in text and "signature" in text_lower and len(text) < 100)):
                # Only insert if paragraph is mostly empty or just contains placeholder text
                if self._is_placeholder(text) or text_lower in ["signature:", "signature"]:
                    # Check if this is user's side (not company side)
                    if not self._is_company_side(text):
                        if self._insert_signature_image(p):
                            fixes += 1
        
        # Scan all tables for signature fields
        for table in doc.tables:
            # Check if this table has a multi-column signature section (like LOA/NDA)
            # Look for headers that indicate left=company, right=user
            has_multi_column_signature = False
            user_column_idx = None  # Track which column is user's side
            
            # Check header rows to identify column structure
            for header_row_idx in range(min(5, len(table.rows))):
                header_row = table.rows[header_row_idx]
                if len(header_row.cells) >= 2:
                    left_header = (header_row.cells[0].text or "").strip().lower()
                    right_header = (header_row.cells[-1].text or "").strip().lower()
                    
                    # Check if left has company name and right has user indicator
                    if (("accolite" in left_header or "company" in left_header) and 
                        ("you" in right_header or (name and name.lower() in right_header))):
                        has_multi_column_signature = True
                        user_column_idx = len(header_row.cells) - 1  # Rightmost column
                        break
            
            for row_idx, row in enumerate(table.rows):
                if len(row.cells) < 2:
                    continue
                
                # Check all cells for signature labels
                has_signature_label = False
                signature_label_col = None
                
                for cell_idx, cell in enumerate(row.cells):
                    cell_text = (cell.text or "").strip().lower()
                    if "signature" in cell_text:
                        has_signature_label = True
                        signature_label_col = cell_idx
                        break
                
                if has_signature_label:
                    # For multi-column signature sections, only fill right side (user's side)
                    if has_multi_column_signature and user_column_idx is not None:
                        # Only fill the rightmost column (user's side), skip left column
                        if len(row.cells) > user_column_idx:
                            user_cell = row.cells[user_column_idx]
                            cell_content = (user_cell.text or "").strip()
                            if self._is_placeholder(cell_content) or not cell_content:
                                if self._insert_signature_in_cell(user_cell):
                                    fixes += 1
                        # Skip left side signature - don't fill column 0 if it's company side
                        continue
                    
                    # For single-column or non-multi-column tables, fill normally
                    # Fill the cell after the signature label, or rightmost if label is in leftmost
                    if signature_label_col is not None and signature_label_col < len(row.cells) - 1:
                        # Fill the cell after the label
                        fill_cell = row.cells[signature_label_col + 1]
                    else:
                        # Fill rightmost cell
                        fill_cell = row.cells[-1]
                    
                    cell_content = (fill_cell.text or "").strip()
                    if self._is_placeholder(cell_content) or not cell_content:
                        if self._insert_signature_in_cell(fill_cell):
                            fixes += 1
                
                # For multi-column signature sections, also check for empty cells after "Print Name" rows
                # These are often signature cells in LOA/NDA forms (empty cells that should have signatures)
                if has_multi_column_signature and user_column_idx is not None:
                    left_text = (row.cells[0].text or "").strip().lower()
                    # Check if this row has "Print Name" in the left cell - signature is usually in the NEXT row
                    if "print name" in left_text:
                        # Check the next row - it might be the signature row (often empty)
                        if row_idx + 1 < len(table.rows):
                            next_row = table.rows[row_idx + 1]
                            if len(next_row.cells) > user_column_idx:
                                # Check if right cell is empty (likely signature cell)
                                user_cell = next_row.cells[user_column_idx]
                                cell_content = (user_cell.text or "").strip()
                                next_left = (next_row.cells[0].text or "").strip().lower()
                                # If next row's left cell is "Signature" or empty, and right cell is empty, it's signature
                                if ((next_left == "signature" or next_left == "" or self._is_placeholder(next_left)) and 
                                    (self._is_placeholder(cell_content) or not cell_content) and 
                                    not self._cell_has_image(user_cell)):
                                    print(f"   Found signature cell in next row after 'Print Name'")
                                    if self._insert_signature_in_cell(user_cell):
                                        fixes += 1
                        # Also check the current row's right cell if it's empty and we're in signature section
                        if len(row.cells) > user_column_idx:
                            user_cell = row.cells[user_column_idx]
                            cell_content = (user_cell.text or "").strip()
                            # If right cell is empty and we just filled "Print Name" on left, this might be signature
                            if (self._is_placeholder(cell_content) or not cell_content) and not self._cell_has_image(user_cell):
                                # Check if previous row was also "Print Name" or similar to confirm we're in signature section
                                if row_idx > 0:
                                    prev_row = table.rows[row_idx - 1]
                                    if len(prev_row.cells) > 0:
                                        prev_left = (prev_row.cells[0].text or "").strip().lower()
                                        if "print name" in prev_left or "title" in prev_left or "date" in prev_left:
                                            # This is likely a signature cell - fill it
                                            print(f"   Found signature cell in current row (after signature section indicators)")
                                            if self._insert_signature_in_cell(user_cell):
                                                fixes += 1
                
                # Check individual cells for standalone signature fields
                for cell_idx, cell in enumerate(row.cells):
                    cell_text = (cell.text or "").strip().lower()
                    if cell_text == "signature:" or cell_text == "signature":
                        # For multi-column tables, only fill user's column
                        if has_multi_column_signature and user_column_idx is not None:
                            if cell_idx != user_column_idx:
                                continue  # Skip company side
                        
                        if self._insert_signature_in_cell(cell):
                            fixes += 1
        
        return fixes

    def populate_form_smart(self, template_path: str, output_path: str) -> bool:
        print(f"🤖 Smart Processing: {os.path.basename(template_path)}")

        structure = self.extract_form_structure(template_path)
        print(f"  📋 Form Type: {structure['form_type']}")

        doc = Document(template_path)
        fixes_applied = 0
        
        # Fill signatures everywhere first (before other field fills)
        fixes_applied += self._fill_signatures_everywhere(doc)

        # Force simple 6-field behavior for forms that just need the basics
        name_lower = os.path.basename(template_path).lower()
        if any(k in name_lower for k in [
            "declarationforpfaccount linking with aadhar",
            "bounteous_hyd_letterhead template_april 2025",
            "declarationformforpfaccountpdf"
        ]):
            fixes_applied += self._fill_simple_6fields_everywhere(doc)
            print(f"  🔧 Applied {fixes_applied} fixes (simple 6-field)")
            doc.save(output_path)
            print(f"  ✅ Saved: {os.path.basename(output_path)}")
            return True

        if structure["form_type"] == "background_verification":
            fixes_applied += self.fill_background_verification_form(doc, structure)
        elif structure["form_type"] == "nda":
            fixes_applied += self.fill_nda_form(doc, structure)
        elif structure["form_type"] == "declaration":
            fixes_applied += self.fill_declaration_form(doc, structure)
            fixes_applied += self._fill_current_address_everywhere(doc)  # ensure address in tables too
        elif structure["form_type"] == "gratuity":
            fixes_applied += self.fill_gratuity_form(doc, structure)
            fixes_applied += self._fill_current_address_everywhere(doc)
        elif structure["form_type"] == "loa":
            fixes_applied += self.fill_loa_form(doc, structure)
            fixes_applied += self._fill_current_address_everywhere(doc)
        elif structure["form_type"] == "pf_account":
            fixes_applied += self.fill_pf_account_form(doc, structure)
            fixes_applied += self._fill_current_address_everywhere(doc)
        elif structure["form_type"] == "epf_nomination":
            print(f"  🎯 Processing EPF Nomination Form")
            print(f"  📷 Signature image path: {self.signature_image_path}")
            fixes_applied += self.fill_epf_nomination_form(doc, structure)
            fixes_applied += self._fill_current_address_everywhere(doc)
        else:
            fixes_applied += self.fill_general_form(doc, structure)
            fixes_applied += self._fill_current_address_everywhere(doc)

        print(f"  🔧 Applied {fixes_applied} fixes")
        doc.save(output_path)
        print(f"  ✅ Saved: {os.path.basename(output_path)}")
        return True

    def _heading_for_table(self, tinfo: Dict) -> str:
        heading = (tinfo.get("heading") or "").strip().lower()
        if not heading or heading in {"personal details:", "employment history:", "education:"}:
            try:
                first_row = tinfo["rows"][0]["cells"]
                probe = " ".join(c["text"] for c in first_row[:2]).strip().lower()
                if any(k in probe for k in ["highest qualification", "previous", "professional"]):
                    heading = probe
            except Exception:
                pass
        heading = heading.replace("details of ", "")
        return heading

    def fill_background_verification_form(self, doc: Document, structure: Dict) -> int:
        fixes_applied = 0
        
        # 0) Fill signature fields in paragraphs (standalone signature fields)
        for p in doc.paragraphs:
            text = (p.text or "").strip().lower()
            # Look for standalone "Signature:" or "Signature" fields
            if text == "signature:" or text == "signature" or (text.endswith("signature:") and len(text) < 50):
                if self._insert_signature_image(p):
                    fixes_applied += 1

        # 1) Paragraph fields that end with ":" → insert values
        for info in structure["paragraphs"]:
            if info["is_field"] and info["text"].strip().endswith(":"):
                ft = info["field_type"]
                
                # Handle signature fields with image insertion
                if ft == "signature":
                    p = doc.paragraphs[info["index"]]
                    if self._insert_signature_image(p):
                        fixes_applied += 1
                    continue
                
                val = self.get_field_value(ft)
                if val:
                    p = doc.paragraphs[info["index"]]
                    original_text = info["text"]
                    
                    # Special handling for expiry date with employment history
                    if ft == "passport_expiry_date_with_employment":
                        # Keep "Employment History:" part intact
                        p.text = f"Expiry Date: {val} Employment History:"
                    else:
                        label = original_text.rstrip(":")
                        p.text = f"{label}: {val}"
                    fixes_applied += 1

        # 2) Prep data
        employment_history = self.emp
        edu_hist = self.edu
        address_history = self.addr
        refs = self.refs
        gaps = self.gaps

        # education list
        edu_list = []
        if edu_hist:
            if "highest_qualification" in edu_hist:
                edu_list.append(edu_hist["highest_qualification"])
            if "previous_qualification" in edu_hist:
                edu_list.append(edu_hist["previous_qualification"])

        # address list in canonical order
        addr_list = []
        # Use address_list if provided (from frontend), otherwise build from address_history
        if "address_list" in self.form_fields and self.form_fields["address_list"]:
            addr_list = self.form_fields["address_list"]
        elif address_history:
            for tag in ["current", "previous", "permanent"]:
                if tag in address_history:
                    dd = dict(address_history[tag])
                    dd["address_type"] = tag
                    addr_list.append(dd)

        # gaps can be dict or list
        gap_list = gaps if isinstance(gaps, list) else ([gaps] if gaps else [])

        counters = {"employment": 0, "education": 0, "address": 0, "reference": 0, "gap": 0, "previous_address": 0}

        # logical education slots for forms that have more boxes than data
        edu_slots = []
        if len(edu_list) >= 1:
            edu_slots.append(edu_list[0])  # highest
        if len(edu_list) >= 2:
            edu_slots.append(edu_list[1])  # previous I
        while len(edu_slots) < 5:
            edu_slots.append(None)

        # 3) Fill signature fields in tables first (before other table fills)
        for tinfo in structure["tables"]:
            table = doc.tables[tinfo["index"]]
            for row in table.rows:
                if len(row.cells) < 2:
                    continue
                # Check if left cell contains "signature" label
                left_text = (row.cells[0].text or "").strip().lower()
                if "signature" in left_text:
                    # Insert signature in the rightmost cell
                    if self._insert_signature_in_cell(row.cells[-1]):
                        fixes_applied += 1
                # Also check all cells for standalone signature fields
                for cell_idx, cell in enumerate(row.cells):
                    cell_text = (cell.text or "").strip().lower()
                    if cell_text == "signature:" or cell_text == "signature" or (cell_text.endswith("signature:") and len(cell_text) < 50):
                        if self._insert_signature_in_cell(cell):
                            fixes_applied += 1
        
        # 4) Fill other tables
        for tinfo in structure["tables"]:
            table = doc.tables[tinfo["index"]]
            section = self._classify_table(table)

            if section == "employment":
                i = counters["employment"]; counters["employment"] += 1
                if i < len(employment_history):
                    fixes_applied += self._fill_employment_table(table, employment_history[i], i)
                else:
                    # Don't fill if we don't have data for this employment table
                    self._clear_right_cells(table)

            elif section == "education":
                # Use a simple approach: fill the first N education tables with available data
                i = counters["education"]
                counters["education"] += 1
                
                # For the first education table, pass the entire list to handle forms with multiple sections in one table
                if i == 0 and len(edu_list) > 0:
                    fixes_applied += self._fill_education_table(table, edu_list, i)
                elif i < len(edu_slots) and edu_slots[i] is not None:
                    fixes_applied += self._fill_education_table(table, edu_slots[i], i)
                else:
                    self._clear_right_cells(table)

            elif section == "address":
                add_fixes, consumed = self._fill_multi_address_table(table, addr_list, counters)
                fixes_applied += add_fixes
                counters["address"] += consumed
                
                # Check if this table also contains reference sections
                if self._looks_like_reference(table) and refs:
                    # Only fill as many references as we have data for
                    fixes_applied += self._fill_reference_table(table, refs, fill_extra_refs=False)
                
                # Check if this table also contains gap sections
                if self._looks_like_gap(table) and gap_list:
                    # Only fill the first gap section, leave the second one empty
                    fixes_applied += self._fill_gap_table(table, gap_list[0], fill_second_gap=False)

            elif section == "reference":
                # Check if this table also contains address sections (mixed table)
                if self._looks_like_address(table) and addr_list:
                    add_fixes, consumed = self._fill_multi_address_table(table, addr_list, counters)
                    fixes_applied += add_fixes
                    counters["address"] += consumed
                
                if refs:
                    # Only fill as many references as we have data for
                    fixes_applied += self._fill_reference_table(table, refs, fill_extra_refs=False)
                else:
                    self._clear_right_cells(table)
                
                # Check if this table also contains gap sections
                if self._looks_like_gap(table) and gap_list:
                    # Only fill the first gap section, leave the second one empty
                    fixes_applied += self._fill_gap_table(table, gap_list[0], fill_second_gap=False)

            elif section == "gap":
                i = counters["gap"]; counters["gap"] += 1
                if i < len(gap_list):
                    fixes_applied += self._fill_gap_table(table, gap_list[i])
                else:
                    self._clear_right_cells(table)

            else:
                # unknown section → leave as-is
                pass

        return fixes_applied

    def _fill_employment_table(self, table, data: Dict, employment_index: int = 0) -> int:
        fixes = [0]
        current_section = None
        sections_filled = 0
        target_section_filled = False  # Track if we've filled the target section
        
        # Check if this table has section headers
        has_section_headers = False
        for row in table.rows:
            left = (row.cells[0].text if row.cells else "").strip().lower()
            if "details of" in left and "employer" in left:
                has_section_headers = True
                break
        
        for row in table.rows:
            left = (row.cells[0].text if row.cells else "").strip().lower()
            
            # Check if this is a new employment section
            if "details of" in left and "employer" in left:
                if "(iv)" in left or "iv)" in left or "iv:" in left:
                    current_section = 3
                elif "(iii)" in left or "iii)" in left or "iii:" in left:
                    current_section = 2
                elif "(ii)" in left or "ii)" in left or "ii:" in left:
                    current_section = 1
                elif "(i)" in left or "i)" in left or "i:" in left:
                    current_section = 0
                else:
                    current_section = 0
                
                # If we encounter a new section after filling our target section, stop
                if target_section_filled:
                    break
                
                # Only fill this section if it matches our employment index
                if current_section != employment_index:
                    continue
                else:
                    # We're in the correct section, start filling
                    target_section_filled = False
            
            # Only fill fields if we're in the correct section OR if there are no section headers
            if current_section == employment_index or (not has_section_headers and employment_index >= 0):
                if any(k in left for k in ["employers name", "employer name", "employers name & branch", "employer name & branch"]):
                    self._fill_cell_if_label(row, ["employer"], data.get("employer_name_and_branch") or data.get("employer_name", ""), fixes)
                if "address" in left:
                    # Try employer address first, then any other address
                    addr_value = data.get("employer_address", "") or data.get("address", "")
                    if addr_value:
                        self._fill_cell_if_label(row, ["address"], addr_value, fixes)
                if "position" in left and ("held" in left or "department" in left or "dept" in left):
                    self._fill_cell_if_label(row, ["position"], data.get("position_and_department") or data.get("position_department", ""), fixes)
                if "employment period" in left:
                    self._fill_cell_if_label(row, ["employment period"], get_nested_value(data, "employment_period"), fixes)
                if "employee code" in left:
                    self._fill_cell_if_label(row, ["employee code"], data.get("employee_code", ""), fixes)
                if "last salary" in left:
                    self._fill_cell_if_label(row, ["salary"], data.get("last_salary", ""), fixes)
                if "reason for leaving" in left:
                    self._fill_cell_if_label(row, ["reason"], data.get("reason_for_leaving", ""), fixes)
                if "reporting manager" in left:
                    self._fill_cell_if_label(row, ["reporting manager"], data.get("reporting_manager", ""), fixes)
                if "telephone" in left or "landline" in left:
                    self._fill_cell_if_label(row, ["telephone"], data.get("landline", ""), fixes)
                if "verify" in left and "employment" in left:
                    self._fill_cell_if_label(row, ["verify"], "Yes" if data.get("can_verify") else "No", fixes)
                if "agency" in left and "details" in left:
                    self._fill_cell_if_label(row, ["agency"], data.get("agency_details") or data.get("contract_agency", ""), fixes)
        return fixes[0]

    def _fill_education_row(self, row, data: Dict) -> int:
        """Helper function to fill a single education row with data."""
        fixes = 0
        if not row.cells or len(row.cells) < 2:
            return 0
            
        left = (row.cells[0].text or "").strip().lower()
        
        if "university" in left and "college" in left:
            # Fill only university name, not the address
            uni = data.get("university_and_college", "")
            if uni and len(row.cells) >= 2:
                row.cells[-1].text = self._blank(uni)
                fixes += 1
        elif "location" in left and ("town" in left or "city" in left or "address" in left):
            # Fill the location/address field separately
            location = data.get("location_full_address", "")
            if location and len(row.cells) >= 2:
                row.cells[-1].text = self._blank(location)
                fixes += 1
        elif "period of the course" in left or "period" in left and "course" in left:
            period_data = data.get("period", {})
            if isinstance(period_data, dict):
                start_date = period_data.get("start", "")
                end_date = period_data.get("end", "")
                if start_date and end_date:
                    # Format dates properly (e.g., "Aug 2022 - May 2024")
                    from datetime import datetime
                    try:
                        start_dt = datetime.strptime(start_date, "%Y-%m-%d")
                        end_dt = datetime.strptime(end_date, "%Y-%m-%d")
                        formatted_period = f"{start_dt.strftime('%b %Y')} - {end_dt.strftime('%b %Y')}"
                        if len(row.cells) >= 2:
                            row.cells[-1].text = self._blank(formatted_period)
                            fixes += 1
                    except:
                        # Fallback to raw period if date parsing fails
                        raw = period_data.get("raw", "")
                        if raw and len(row.cells) >= 2:
                            row.cells[-1].text = self._blank(raw)
                            fixes += 1
                else:
                    raw = period_data.get("raw", "")
                    if raw and len(row.cells) >= 2:
                        row.cells[-1].text = self._blank(raw)
                        fixes += 1
        elif any(k in left for k in ["degree", "diploma", "course"]):
            degree = data.get("degree_or_course", "")
            if degree and len(row.cells) >= 2:
                row.cells[-1].text = self._blank(degree)
                fixes += 1
        elif any(k in left for k in ["roll", "registration", "seat"]):
            roll = data.get("roll_or_registration", "")
            if roll and len(row.cells) >= 2:
                row.cells[-1].text = self._blank(roll)
                fixes += 1
        
        return fixes

    def _fill_education_table(self, table, data: Dict, education_index: int = 0) -> int:
        """
        Fill education table. If data is a dict, fill one section at education_index.
        If data is a list, fill multiple sections (for forms with multiple sections in one table).
        """
        # If data is a list, fill all sections in the table
        if isinstance(data, list):
            fixes = 0
            edu_data_list = data
            sections_encountered = 0
            current_section = None
            current_data = None
            
            for row in table.rows:
                left = (row.cells[0].text if row.cells else "").strip().lower()
                
                # Check if this is a new education section header
                if "details of" in left and ("qualification" in left or "professional" in left):
                    # Skip "Professional Qualifications or Certifications" sections for regular education data
                    if "professional" in left and "certification" in left:
                        current_section = None
                        current_data = None
                        continue
                    
                    # Activate this section if we have data for it
                    if sections_encountered < len(edu_data_list):
                        current_section = "active"
                        current_data = edu_data_list[sections_encountered]
                    else:
                        current_section = None
                        current_data = None
                    sections_encountered += 1
                
                # Fill fields if we're in an active section with data
                if current_section == "active" and current_data:
                    fixes += self._fill_education_row(row, current_data)
            
            return fixes
        
        # Original logic for single education entry (dict)
        fixes = [0]
        current_section = None
        sections_encountered = 0
        
        for row in table.rows:
            left = (row.cells[0].text if row.cells else "").strip().lower()
            
            # Check if this is a new education section header
            if "details of" in left and ("qualification" in left or "professional" in left):
                # Skip "Professional Qualifications or Certifications" sections for regular education data
                if "professional" in left and "certification" in left:
                    current_section = None  # Skip this section
                    continue
                
                # Count the regular education sections and activate only the one matching our education_index
                # education_index: 0 = first section (Highest), 1 = second section (Previous I), etc.
                if sections_encountered == education_index:
                    current_section = "active"
                else:
                    current_section = None
                sections_encountered += 1
            
            # Only fill fields if we're in an active section
            if current_section == "active":
                if "university" in left and "college" in left:
                    # Fill only university name, not the address
                    uni = data.get("university_and_college", "")
                    self._fill_cell_if_label(row, ["university", "college"], uni, fixes)
                elif "location" in left and ("town" in left or "city" in left or "address" in left):
                    # Fill the location/address field separately
                    location = data.get("location_full_address", "")
                    self._fill_cell_if_label(row, ["location", "address"], location, fixes)
                elif any(k in left for k in ["degree", "diploma", "course"]):
                    self._fill_cell_if_label(row, ["degree"], data.get("degree_or_course", ""), fixes)
                elif "period of the course" in left:
                    period_data = data.get("period", {})
                    if isinstance(period_data, dict):
                        start_date = period_data.get("start", "")
                        end_date = period_data.get("end", "")
                        if start_date and end_date:
                            # Format dates properly (e.g., "Aug 2022 - May 2024")
                            from datetime import datetime
                            try:
                                start_dt = datetime.strptime(start_date, "%Y-%m-%d")
                                end_dt = datetime.strptime(end_date, "%Y-%m-%d")
                                formatted_period = f"{start_dt.strftime('%b %Y')} - {end_dt.strftime('%b %Y')}"
                                self._fill_cell_if_label(row, ["period"], formatted_period, fixes)
                            except:
                                # Fallback to raw period if date parsing fails
                                self._fill_cell_if_label(row, ["period"], period_data.get("raw", ""), fixes)
                        else:
                            self._fill_cell_if_label(row, ["period"], period_data.get("raw", ""), fixes)
                    else:
                        self._fill_cell_if_label(row, ["period"], str(period_data), fixes)
                elif any(k in left for k in ["roll", "registration", "seat"]):
                    self._fill_cell_if_label(row, ["roll"], data.get("roll_or_registration", ""), fixes)
        return fixes[0]

    def _fill_multi_address_table(self, table, address_list: List[Dict], counters: Dict) -> Tuple[int, int]:
        """Fill a table that contains multiple address sections.
        Returns (fixes_applied, addresses_consumed)."""
        fixes = 0
        idx = counters["address"]  # Use the global address counter
        last_used = None
        consumed_for_this_table = 0
        in_section = False
        # Use global previous_address counter that persists across tables
        previous_address_counter = counters["previous_address"]

        for row in table.rows:
            if len(row.cells) < 2:
                continue
            left = (row.cells[0].text or "").strip().lower()
            right = row.cells[-1]  # ALWAYS last cell

            if "complete" in left and "address" in left:
                in_section = True
                last_used = None
                
                # Smart matching: match address type with section label
                target_address = None
                
                if "permanent" in left:
                    # Look for permanent address in the list
                    target_address = next((addr for addr in address_list if addr.get("address_type") == "permanent"), None)
                elif "previous" in left:
                    # Look for previous address(es) in the list
                    # For "Previous Address 1", "Previous Address 2", etc., fill only if we have that many previous addresses
                    previous_addresses = [addr for addr in address_list if addr.get("address_type") == "previous"]
                    if previous_address_counter < len(previous_addresses):
                        target_address = previous_addresses[previous_address_counter]
                    previous_address_counter += 1
                elif "current" in left:
                    # Look for current address in the list
                    target_address = next((addr for addr in address_list if addr.get("address_type") == "current"), None)
                else:
                    # No specific type mentioned, use sequential logic
                    if idx < len(address_list):
                        target_address = address_list[idx]
                    idx += 1
                    consumed_for_this_table += 1
                
                if target_address:
                    last_used = target_address
                    # Fill complete address with the full address
                    right.text = self._blank(target_address.get("town_or_city_name"))
                    fixes += 1
                else:
                    right.text = ""
                continue

            if in_section and last_used is None:
                if ("town" in left or "city" in left or
                    "duration of stay" in left or "phone" in left):
                    right.text = ""
                continue

            if last_used:
                if "town" in left or "city" in left:
                    # Extract city name from the full address
                    full_address = last_used.get("town_or_city_name", "")
                    city_name = self._extract_city_name(full_address)
                    right.text = self._blank(city_name); fixes += 1
                elif "duration of stay" in left:
                    dur = last_used.get("duration_of_stay", {})
                    val = dur.get("raw") if isinstance(dur, dict) else (str(dur) if dur else "")
                    right.text = self._blank(val); fixes += 1
                elif "phone" in left:
                    right.text = self._blank(last_used.get("phone_number")); fixes += 1

        # Update the global previous_address counter for next table
        counters["previous_address"] = previous_address_counter
        return fixes, consumed_for_this_table

    # ------- Reference & Gap (now using _row_label) -------
    def _fill_reference_table(self, table, data: Union[Dict, List[Dict]], fill_extra_refs: bool = True) -> int:
        refs = data if isinstance(data, list) else [data]
        fixes = 0
        ref_idx = 0

        def set_last(row, v):
            nonlocal fixes
            if len(row.cells) >= 2:
                row.cells[-1].text = self._blank(v)
                fixes += 1

        # detect headings like "Reference 1", "Reference (i)", etc. on ANY row
        have_named_slots = any(
            re.search(r"\breference\s*(1|\(i\)|i|2|\(ii\)|ii)\b", self._row_label(row))
            for row in table.rows
        )

        if have_named_slots:
            in_reference_section = False
            for row in table.rows:
                blob = self._row_label(row)

                if re.search(r"\breference\s*(1|\(i\)|i)\b", blob):
                    ref_idx = 0
                    in_reference_section = True
                    continue
                if re.search(r"\breference\s*(2|\(ii\)|ii)\b", blob):
                    ref_idx = 1
                    in_reference_section = True
                    continue
                
                # Check if we're leaving the reference section (e.g., gap section starts)
                if "gap" in blob.lower() or "details of gap" in blob.lower() or "details of gap if any" in blob.lower():
                    in_reference_section = False
                    continue

                # Only process rows that are in the reference section
                if not in_reference_section:
                    continue

                left = blob  # whole-row label
                ref = refs[ref_idx] if ref_idx < len(refs) else None

                # If we don't have data for this reference and fill_extra_refs is False, skip it
                if ref is None or (ref_idx >= len(refs) and not fill_extra_refs):
                    if any(k in left for k in ["name", "ph", "phone", "mobile", "designation", "company", "organization", "org", "email", "relationship", "years"]):
                        if len(row.cells) >= 2:
                            row.cells[-1].text = ""
                    continue

                if "designation" in left or "company" in left or "organization" in left or "org" in left:
                    set_last(row, ref.get("designation_and_company", ""))
                elif "ph" in left or "phone" in left or "mobile" in left or "contact" in left:
                    set_last(row, ref.get("phone", ""))
                elif "name" in left:
                    set_last(row, ref.get("name", ""))
                elif "email" in left or "mail" in left:
                    set_last(row, ref.get("email", ref.get("phone", "")))
                elif "relationship" in left:
                    set_last(row, ref.get("relationship", ""))
                elif "years known" in left or "known for" in left or "years" in left:
                    set_last(row, ref.get("years_known", ""))

        else:
            # sequential groups: Name → Phone → Designation → Email → Relationship → Years … then next ref
            fields_written = set()
            for row in table.rows:
                left = self._row_label(row)
                ref = refs[ref_idx] if ref_idx < len(refs) else None

                # If we don't have data for this reference and fill_extra_refs is False, skip it
                if ref is None or (ref_idx >= len(refs) and not fill_extra_refs):
                    if any(k in left for k in ["name", "ph", "phone", "mobile", "designation", "company", "organization", "org", "email", "relationship", "years"]):
                        if len(row.cells) >= 2:
                            row.cells[-1].text = ""
                    continue

                if "designation" in left or "company" in left or "organization" in left or "org" in left:
                    set_last(row, ref.get("designation_and_company", "")); fields_written.add("designation")
                elif "ph" in left or "phone" in left or "mobile" in left or "contact" in left:
                    set_last(row, ref.get("phone", "")); fields_written.add("phone")
                elif "name" in left:
                    if "name" in fields_written:
                        ref_idx += 1
                        fields_written = set()
                        ref = refs[ref_idx] if ref_idx < len(refs) else None
                        if ref is None:
                            if len(row.cells) >= 2:
                                row.cells[-1].text = ""
                            continue
                    set_last(row, ref.get("name", "")); fields_written.add("name")
                elif "email" in left or "mail" in left:
                    set_last(row, ref.get("email", ref.get("phone", ""))); fields_written.add("email")
                elif "relationship" in left:
                    set_last(row, ref.get("relationship", "")); fields_written.add("relationship")
                elif "years known" in left or "known for" in left or "years" in left:
                    set_last(row, ref.get("years_known", "")); fields_written.add("years")

        return fixes

    def _fill_gap_table(self, table, data: Dict, fill_second_gap: bool = True) -> int:
        fixes = [0]
        gap_section = 1  # Track which gap section we're filling
        
        for row in table.rows:
            left = self._row_label(row)
            
            # Check if we're in the second gap section
            if "gap 2" in left.lower() or "reason for gap 2" in left.lower():
                gap_section = 2
                if not fill_second_gap:
                    continue  # Skip filling second gap section
            
            if ("reason for gap" in left) or (left.startswith("reason")) or ("gap reason" in left):
                if gap_section == 1 or fill_second_gap:
                    self._fill_cell_if_label(row, ["reason"], data.get("reason", ""), fixes)
            elif ("period of gap" in left) or ("period" in left) or ("from" in left) or ("to" in left) or ("duration" in left):
                if gap_section == 1 or fill_second_gap:
                    self._fill_cell_if_label(row, ["period"], get_nested_value(data, "period"), fixes)
            elif ("address stayed during the gap" in left) or ("address during gap" in left) or ("address stayed" in left):
                if gap_section == 1 or fill_second_gap:
                    self._fill_cell_if_label(row, ["address"], data.get("address_during_gap", ""), fixes)
        return fixes[0]

    # ----- Other forms -----
    def fill_nda_form(self, doc: Document, structure: Dict) -> int:
        fixes = 0
        name = self._get("name")
        title = self.emp[0].get("position_and_department", "") if self.emp else ""

        for p in doc.paragraphs:
            t = p.text.strip()
            if re.search(r"\bPrint Name\b.*:\s*$", t):
                p.text = f"Print Name: {name}"; fixes += 1
            elif re.fullmatch(r"\s*Title\s*:\s*", t):
                p.text = f"Title: {title}"; fixes += 1
            elif re.fullmatch(r"\s*Date\s*:\s*", t):
                p.text = f"Date: {today_str()}"; fixes += 1

        # Find signature cells in tables - look for empty cells after signature labels
        for table in doc.tables:
            # Check if this is a multi-column signature table (company vs user)
            has_multi_column = False
            user_column_idx = None
            
            # Check header rows to identify structure
            for header_row_idx in range(min(5, len(table.rows))):
                header_row = table.rows[header_row_idx]
                if len(header_row.cells) >= 2:
                    left_header = (header_row.cells[0].text or "").strip().lower()
                    right_header = (header_row.cells[-1].text or "").strip().lower()
                    if ("accolite" in left_header or "company" in left_header) and ("you" in right_header or (name and name.lower() in right_header)):
                        has_multi_column = True
                        user_column_idx = len(header_row.cells) - 1
                        break
            
            for row_idx, row in enumerate(table.rows):
                if len(row.cells) < 2:
                    continue
                
                # Check if left cell contains signature label
                L = (row.cells[0].text or "").strip()
                left_text = L.lower()
                
                # IMPORTANT: Check for signature FIRST before filling other fields
                # This prevents text from overwriting the signature image
                if "signature" in left_text:
                    # This is a signature row - insert signature in the right column only
                    print(f"🔍 NDA: Found signature row, left='{L}', has_multi_column={has_multi_column}, user_column_idx={user_column_idx}")
                    if has_multi_column and user_column_idx is not None and len(row.cells) > user_column_idx:
                        # Only fill user's side (rightmost column)
                        user_cell = row.cells[user_column_idx]
                        cell_content = (user_cell.text or "").strip()
                        print(f"   NDA: User cell content='{cell_content}', is_placeholder={self._is_placeholder(cell_content)}")
                        # Insert signature if cell is empty, placeholder, or contains "Signature" text
                        if (self._is_placeholder(cell_content) or not cell_content or 
                            cell_content.lower() == "signature"):
                            print(f"   NDA: Attempting signature insertion...")
                            if self._insert_signature_in_cell(user_cell):
                                fixes += 1
                                print(f"   NDA: ✅ Signature inserted!")
                            else:
                                print(f"   NDA: ❌ Signature insertion failed")
                    elif not has_multi_column:
                        # Single column - fill rightmost cell
                        cell_content = (row.cells[-1].text or "").strip()
                        if (self._is_placeholder(cell_content) or not cell_content or 
                            cell_content.lower() == "signature"):
                            if self._insert_signature_in_cell(row.cells[-1]):
                                fixes += 1
                    continue  # Skip to next row, don't process this as a text field
                
                # Process other text fields only if not a signature row
                if re.search(r"\bPrint Name\b", L, re.I):
                    # Check if cell already has an image (signature), don't overwrite
                    if not self._cell_has_image(row.cells[-1]):
                        row.cells[-1].text = name; fixes += 1
                elif re.search(r"\bTitle\b", L, re.I):
                    # Check if cell already has an image (signature), don't overwrite
                    if not self._cell_has_image(row.cells[-1]):
                        row.cells[-1].text = title; fixes += 1
                elif re.search(r"\bDate\b", L, re.I):
                    # Check if cell already has an image (signature), don't overwrite
                    if not self._cell_has_image(row.cells[-1]):
                        row.cells[-1].text = today_str(); fixes += 1
                # Fill "Print Name" fields in the signature section
                elif L == "Print Name":
                    # Check if cell already has an image (signature), don't overwrite
                    if not self._cell_has_image(row.cells[-1]):
                        row.cells[-1].text = name; fixes += 1
                
                # After processing all text fields, check if we're in a signature section and fill empty signature cells
                # This handles cases where signature cells don't have explicit "Signature" labels
                if has_multi_column and user_column_idx is not None:
                    # If we're in a row that has "Print Name", "Title", or "Date" in left column,
                    # and the next row has empty cells, it might be the signature row
                    left_text = (row.cells[0].text or "").strip().lower()
                    if any(label in left_text for label in ["print name", "title", "date"]):
                        # Check if next row has empty right cell (likely signature)
                        if row_idx + 1 < len(table.rows):
                            next_row = table.rows[row_idx + 1]
                            if len(next_row.cells) > user_column_idx:
                                next_left = (next_row.cells[0].text or "").strip().lower()
                                next_right = (next_row.cells[user_column_idx].text or "").strip()
                                # If next row's left is "Signature" or empty, and right is empty, fill signature
                                if (next_left == "signature" or next_left == "" or self._is_placeholder(next_left)) and \
                                   (self._is_placeholder(next_right) or not next_right) and \
                                   not self._cell_has_image(next_row.cells[user_column_idx]):
                                    print(f"   NDA: Found empty signature cell after '{L}' row")
                                    if self._insert_signature_in_cell(next_row.cells[user_column_idx]):
                                        fixes += 1
                
                # Check for "By:" fields in any cell of the row
                if len(row.cells) >= 3:
                    # Look for the pattern: "By:" in Cell 0, empty Cell 1, "By:" in Cell 2
                    if (row.cells[0].text.strip() == "By:" and 
                        row.cells[1].text.strip() == "" and 
                        row.cells[2].text.strip() == "By:"):
                        # Fill Cell 2 (the rightmost "By:" field) with employee name
                        # But check if there's already a signature image
                        if not self._cell_has_image(row.cells[2]):
                            row.cells[2].text = f"By: {name}"; fixes += 1
        return fixes

    def fill_declaration_form(self, doc: Document, structure: Dict) -> int:
        """Simplified declaration form filling - fills data as text in paragraphs and tables."""
        return self._fill_declaration_form(doc)

    def fill_gratuity_form(self, doc: Document, structure: Dict) -> int:
        fixes = 0
        name = self._get("name")
        father = self._get("father_name", "fathers_name")
        title = self.emp[0].get("position_and_department", "") if self.emp else ""
        perm = (self.addr.get("permanent") or self.addr.get("current") or {})
        paddr = perm.get("town_or_city_name", "")

        for p in doc.paragraphs:
            t = p.text
            if re.search(r"Name of employee in full", t, re.I):
                p.text = f"Name of employee in full: {name}"; fixes += 1
            elif re.search(r"Father.*Name", t, re.I):
                p.text = f"{t.strip()} {father}"; fixes += 1
            elif re.search(r"Post held|Designation|Position", t, re.I):
                p.text = f"{t.strip()} {title}"; fixes += 1
            elif re.search(r"Permanent address", t, re.I):
                p.text = f"Permanent address: {paddr}"; fixes += 1
            elif re.fullmatch(r"\s*Date\s*:\s*", t):
                p.text = f"Date: {today_str()}"; fixes += 1
        return fixes

    def fill_loa_form(self, doc: Document, structure: Dict) -> int:
        fixes = 0
        name = self._get("name")
        title = self.emp[0].get("position_and_department", "") if self.emp else ""
        employer = self.emp[0].get("employer_name_and_branch", self.emp[0].get("employer_name", "")) if self.emp else ""

        for p in doc.paragraphs:
            t = p.text.strip()
            if re.fullmatch(r"\s*Name\s*:\s*", t):
                p.text = f"Name: {name}"; fixes += 1
            elif re.fullmatch(r"\s*Position\s*:\s*", t):
                p.text = f"Position: {title}"; fixes += 1
            elif re.fullmatch(r"\s*Employer\s*:\s*", t):
                p.text = f"Employer: {employer}"; fixes += 1
            elif re.fullmatch(r"\s*Date\s*:\s*", t):
                p.text = f"Date: {today_str()}"; fixes += 1

        # Find signature cells in tables - look for empty cells after signature labels
        for table in doc.tables:
            # Check if this is a multi-column signature table (company vs user)
            has_multi_column = False
            user_column_idx = None
            
            # Check header rows to identify structure
            for header_row_idx in range(min(5, len(table.rows))):
                header_row = table.rows[header_row_idx]
                if len(header_row.cells) >= 2:
                    left_header = (header_row.cells[0].text or "").strip().lower()
                    right_header = (header_row.cells[-1].text or "").strip().lower()
                    if ("accolite" in left_header or "company" in left_header) and ("you" in right_header or (name and name.lower() in right_header)):
                        has_multi_column = True
                        user_column_idx = len(header_row.cells) - 1
                        break
            
            for row_idx, row in enumerate(table.rows):
                if len(row.cells) >= 2:
                    L = (row.cells[0].text or "").strip()
                    left_text = L.lower()
                    
                    # IMPORTANT: Check for signature FIRST before filling other fields
                    # This prevents text from overwriting the signature image
                    if "signature" in left_text:
                        # This is a signature row - insert signature in the right column only
                        print(f"🔍 LOA: Found signature row, left='{L}', has_multi_column={has_multi_column}, user_column_idx={user_column_idx}")
                        if has_multi_column and user_column_idx is not None and len(row.cells) > user_column_idx:
                            # Only fill user's side (rightmost column)
                            user_cell = row.cells[user_column_idx]
                            cell_content = (user_cell.text or "").strip()
                            print(f"   LOA: User cell content='{cell_content}', is_placeholder={self._is_placeholder(cell_content)}")
                            # Insert signature if cell is empty, placeholder, or contains "Signature" text
                            if (self._is_placeholder(cell_content) or not cell_content or 
                                cell_content.lower() == "signature"):
                                print(f"   LOA: Attempting signature insertion...")
                                if self._insert_signature_in_cell(user_cell):
                                    fixes += 1
                                    print(f"   LOA: ✅ Signature inserted!")
                                else:
                                    print(f"   LOA: ❌ Signature insertion failed")
                        elif not has_multi_column:
                            # Single column - fill rightmost cell
                            if self._is_placeholder(row.cells[-1].text) or not row.cells[-1].text.strip():
                                if self._insert_signature_in_cell(row.cells[-1]):
                                    fixes += 1
                        continue  # Skip to next row, don't process this as a text field
                    
                    # Process other text fields only if not a signature row
                    if re.search(r"\bName\b", L, re.I):
                        # Check if cell already has an image (signature), don't overwrite
                        if not self._cell_has_image(row.cells[-1]):
                            row.cells[-1].text = name; fixes += 1
                    elif re.search(r"\bPosition\b", L, re.I):
                        # Check if cell already has an image (signature), don't overwrite
                        if not self._cell_has_image(row.cells[-1]):
                            row.cells[-1].text = title; fixes += 1
                    elif re.search(r"\bEmployer\b", L, re.I):
                        # Check if cell already has an image (signature), don't overwrite
                        if not self._cell_has_image(row.cells[-1]):
                            row.cells[-1].text = employer; fixes += 1
                    elif re.search(r"\bDate\b", L, re.I):
                        # Check if cell already has an image (signature), don't overwrite
                        if not self._cell_has_image(row.cells[-1]):
                            row.cells[-1].text = today_str(); fixes += 1
                    # Fill "Print Name" fields in the signature section
                    elif L == "Print Name":
                        # Check if cell already has an image (signature), don't overwrite
                        if not self._cell_has_image(row.cells[-1]):
                            row.cells[-1].text = name; fixes += 1
                
                # After processing all text fields, check if we're in a signature section and fill empty signature cells
                # This handles cases where signature cells don't have explicit "Signature" labels
                if has_multi_column and user_column_idx is not None:
                    # If we're in a row that has "Print Name", "Title", or "Date" in left column,
                    # and the next row has empty cells, it might be the signature row
                    left_text = (row.cells[0].text or "").strip().lower()
                    if any(label in left_text for label in ["print name", "title", "date", "name", "position"]):
                        # Check if next row has empty right cell (likely signature)
                        if row_idx + 1 < len(table.rows):
                            next_row = table.rows[row_idx + 1]
                            if len(next_row.cells) > user_column_idx:
                                next_left = (next_row.cells[0].text or "").strip().lower()
                                next_right = (next_row.cells[user_column_idx].text or "").strip()
                                # If next row's left is "Signature" or empty, and right is empty, fill signature
                                if (next_left == "signature" or next_left == "" or self._is_placeholder(next_left)) and \
                                   (self._is_placeholder(next_right) or not next_right) and \
                                   not self._cell_has_image(next_row.cells[user_column_idx]):
                                    print(f"   LOA: Found empty signature cell after '{L}' row")
                                    if self._insert_signature_in_cell(next_row.cells[user_column_idx]):
                                        fixes += 1
                
                # Check for "By:" fields in any cell of the row
                if len(row.cells) >= 3:
                    # Look for the pattern: "By:" in Cell 0, empty Cell 1, "By:" in Cell 2
                    if (row.cells[0].text.strip() == "By:" and 
                        row.cells[1].text.strip() == "" and 
                        row.cells[2].text.strip() == "By:"):
                        # Fill Cell 2 (the rightmost "By:" field) with employee name
                        # But check if there's already a signature image
                        if not self._cell_has_image(row.cells[2]):
                            row.cells[2].text = f"By: {name}"; fixes += 1
        return fixes

    def fill_pf_account_form(self, doc: Document, structure: Dict) -> int:
        fixes = 0
        name = self._get("name")
        father = self._get("father_name", "fathers_name")
        email = self._get("email")

        for p in doc.paragraphs:
            t = p.text
            if "I, ………………………………. s/o or d/o ……………………………………...…." in t:
                p.text = t.replace("I, ………………………………. s/o or d/o ……………………………………...….",
                                   f"I, {name} s/o or d/o {father}")
                fixes += 1
            elif re.fullmatch(r"\s*Name:\s*", t):
                p.text = f"Name: {name}"; fixes += 1
            elif re.search(r"Personal\s+Email\s+id\s*:\s*Signature\s*:\s*$", t, re.I):
                p.text = f"Personal Email id: {email}   Signature:"
                # Add signature image if available
                if self.signature_image_path and os.path.exists(self.signature_image_path):
                    try:
                        run = p.add_run()
                        run.add_picture(self.signature_image_path, width=Inches(2.0))
                    except:
                        pass
                fixes += 1
            elif re.fullmatch(r"\s*Date\s*:\s*", t):
                p.text = f"Date: {today_str()}"; fixes += 1
        return fixes

    # ------- Simple 6-field filler for the three “basic” forms -------
    def _fill_simple_6fields_everywhere(self, doc: Document) -> int:
        """
        Fill Name, Father's Name, Email, Address (default to CURRENT), Date
        in both paragraphs and tables. Very permissive matching.
        """
        fixes = 0
        name   = self._get("name")
        father = self._get("father_name", "fathers_name")
        email  = self._get("email")
        cur    = self._current_address_value()
        perm   = self._permanent_address_value()
        prev   = self._previous_address_value()
        today  = today_str()

        # paragraphs
        # paragraphs
        for p in doc.paragraphs:
            txt = p.text or ""
            label_only = txt.split(":", 1)[0].strip() if ":" in txt else txt.strip()
            
            txt = p.text or ""


            if self.DECLARATION_NAME_LINE_RE.search(txt) and self.REL_MARKER_RE.search(txt):
                name = self._get("name")
                father = self._get("father_name", "fathers_name")
                if name and father:
                    new_line = f"I, {name} s/o or d/o {father}"
                    
                    # Replace only textual content inside runs, preserve paragraph structure
                    for run in p.runs:
                        if self.REL_MARKER_RE.search(run.text) or "I" in run.text:
                            run.text = new_line
                            break
                continue


            if self.NAME_LABEL_RE.search(txt):
                p.text = f"{label_only}: {name}"; fixes += 1
                continue

            if self.FATHER_LABEL_RE.search(txt):
                p.text = f"{label_only}: {father}"; fixes += 1
                continue

            # Special handling for "Personal Email id: Signature:" pattern
            if re.search(r"Personal\s+Email\s+id\s*:\s*Signature\s*:\s*$", txt, re.I):
                # Fill email first
                p.text = f"Personal Email id: {email}   Signature:"
                # Then add signature image after the text if available
                if self.signature_image_path and os.path.exists(self.signature_image_path):
                    try:
                        run = p.add_run()
                        run.add_picture(self.signature_image_path, width=Inches(2.0))
                    except:
                        pass
                fixes += 1
                continue

            if self.EMAIL_LABEL_RE.search(txt):
                p.text = f"{label_only}: {email}"; fixes += 1
                continue

            if self.ADDRESS_LABEL_RE.search(txt):
                lab = label_only.lower()
                if "permanent" in lab:
                    addr_val = perm
                elif "previous" in lab:
                    addr_val = prev
                else:
                    addr_val = cur
                p.text = f"{label_only}: {addr_val}"; fixes += 1
                continue

            if self.DATE_LABEL_RE.search(txt):
                p.text = f"Date: {today}"; fixes += 1
                continue

        # tables
        for table in doc.tables:
            for row in table.rows:
                if len(row.cells) < 2: 
                    continue
                left_text = (row.cells[0].text or "")
                last = row.cells[-1]

                # Check for signature fields in tables
                if "signature" in left_text.lower():
                    if self._insert_signature_in_cell(last):
                        fixes += 1
                elif self.NAME_LABEL_RE.search(left_text):
                    last.text = name; fixes += 1
                elif self.FATHER_LABEL_RE.search(left_text):
                    last.text = father; fixes += 1
                elif self.EMAIL_LABEL_RE.search(left_text):
                    last.text = email; fixes += 1
                elif self.ADDRESS_LABEL_RE.search(left_text):
                    l = left_text.lower()
                    if "permanent" in l:
                        last.text = perm
                    elif "previous" in l:
                        last.text = prev
                    else:
                        last.text = cur
                    fixes += 1
                elif self.DATE_LABEL_RE.search(left_text):
                    last.text = today; fixes += 1

        return fixes

    # ------- Fill current address everywhere for non-BGV forms -------
    def _fill_current_address_everywhere(self, doc: Document) -> int:
        """
        Non-BGV rule: fill current by default; honor explicit Permanent/Previous if present.
        Works in both paragraphs and tables.
        """
        fixes = 0
        cur  = self._current_address_value()
        perm = self._permanent_address_value()
        prev = self._previous_address_value()

        if not (cur or perm or prev):
            return 0

        # paragraphs
        for p in doc.paragraphs:
            text = p.text or ""
            if not self.ADDRESS_LABEL_RE.search(text):
                continue
            label = text.split(":", 1)[0] if ":" in text else text
            l = label.lower()
            if "permanent" in l:
                v = perm
            elif "previous" in l:
                v = prev
            else:
                v = cur
            if not v:
                continue
            if ":" in text:
                p.text = f"{label.strip()}: {v}"
            else:
                p.text = f"{label.strip()}: {v}"
            fixes += 1

        # tables
        for table in doc.tables:
            for row in table.rows:
                if not row.cells:
                    continue
                label = row.cells[0].text or ""
                if not self.ADDRESS_LABEL_RE.search(label):
                    continue
                if len(row.cells) < 2:
                    continue
                
                # Skip header rows - they often contain descriptive text like "full address of nominee"
                # Real address labels are typically short (e.g., "Address:", "Current Address:")
                # Header cells tend to be longer and contain "of" or "the"
                if len(label) > 60 or " of " in label.lower() or " the " in label.lower():
                    continue
                
                l = label.lower()
                if "permanent" in l:
                    v = perm
                elif "previous" in l:
                    v = prev
                else:
                    v = cur
                row.cells[-1].text = v
                fixes += 1

        return fixes

    def fill_gratuity_form(self, doc: Document, structure: Dict) -> int:
        """Fill Gratuity form with employee and nominee details."""
        fixes_applied = 0
        personal_details = self.form_fields.get('personal_details', {})
        epf_gratuity = self.form_fields.get('epf_and_gratuity', {})
        
        # Employee details
        employee_name = personal_details.get('name', '')
        gender = personal_details.get('gender', '')
        employee_address = self.addr.get('current', {}).get('town_or_city_name', '') or self.addr.get('permanent', {}).get('town_or_city_name', '')
        
        # Employment details
        post_held = epf_gratuity.get('post_held', '')
        department = epf_gratuity.get('department', '')
        date_of_appointment = epf_gratuity.get('date_of_appointment', {})
        appointment_date = date_of_appointment.get('value', '') if isinstance(date_of_appointment, dict) else str(date_of_appointment)
        religion = epf_gratuity.get('religion', '')
        marital_status = epf_gratuity.get('marital_status', '')
        
        # Nominee details
        nominee = epf_gratuity.get('nominee', {})
        nominee_name = nominee.get('name', '')
        nominee_relationship = nominee.get('relationship', '')
        nominee_address = nominee.get('address', '')
        nominee_dob = nominee.get('date_of_birth', {})
        nominee_dob_value = nominee_dob.get('value', '') if isinstance(nominee_dob, dict) else str(nominee_dob)
        nominee_share = nominee.get('share', '')
        
        # Calculate age from date of birth
        nominee_age = ''
        if nominee_dob_value:
            try:
                from datetime import datetime
                # Try different date formats
                for fmt in ['%Y-%m-%d', '%d-%m-%Y', '%m/%d/%Y', '%d/%m/%Y']:
                    try:
                        dob_date = datetime.strptime(nominee_dob_value, fmt)
                        current_year = datetime.now().year
                        nominee_age = str(current_year - dob_date.year)
                        break
                    except ValueError:
                        continue
            except:
                pass
        
        # Witnesses
        witnesses = epf_gratuity.get('witnesses', [])
        witness_1 = witnesses[0] if len(witnesses) > 0 else ''
        witness_2 = witnesses[1] if len(witnesses) > 1 else ''
        
        # Form details - use today's date
        from datetime import datetime
        today = datetime.now()
        sign_date = today.strftime('%d-%m-%Y')
        sign_place = epf_gratuity.get('form_sign_place', '')
        
        # Fill paragraphs
        for paragraph in doc.paragraphs:
            text = paragraph.text
            # shri.shrimati/Kumari…………………
            if 'shri.shrimati/Kumari' in text and '…' in text:
                gender_prefix = 'Shrimati' if gender.lower() in ['female', 'f'] else 'Shri'
                paragraph.text = text.replace('shri.shrimati/Kumari…………………', f'{gender_prefix} {employee_name}')
                fixes_applied += 1
            # Place………………
            elif text.strip() == 'Place………………':
                paragraph.text = f"Place {sign_place}"
                fixes_applied += 1
            # Date……………….\tSignature/Thumb-impression
            elif 'Date……………….' in text and 'Signature/Thumb-impression' in text:
                paragraph.text = f"Date {sign_date}\tSignature/Thumb-impression"
                fixes_applied += 1
            # Place : Bangalore Date :
            elif 'Place :' in text and 'Date :' in text:
                paragraph.text = f"Place : {sign_place} Date : {sign_date}"
                fixes_applied += 1
            # Name in full and full address of: (witness section)
            elif 'Name in full and full address of:' in text and text.strip() == 'Name in full and full address of:':
                # Format witness details with names and addresses
                witness_details = []
                for witness_name in witnesses:
                    # Use employee address as witness address (since witness addresses not available)
                    witness_details.append(f"{witness_name}, {employee_address}")
                paragraph.text = f"Name in full and full address of: {'; '.join(witness_details)}"
                fixes_applied += 1

        # Fill tables
        for table in doc.tables:
            for row_idx, row in enumerate(table.rows):
                if len(row.cells) >= 4:  # Nominee table
                    # Check if this is the nominee table (first row has headers)
                    if row_idx == 0 and 'Name in full with full address of nominee' in str([cell.text for cell in row.cells]):
                        continue  # Skip header row
                    elif row_idx == 2:  # Second data row (after header row 0 and column number row 1)
                        # Only fill if not already filled
                        if not row.cells[0].text.strip() or row.cells[0].text.strip() == '':
                            row.cells[0].text = f"{nominee_name}, {nominee_address}"  # Name with address
                            row.cells[1].text = nominee_relationship  # Relationship
                            row.cells[2].text = nominee_age  # Age
                            row.cells[3].text = nominee_share  # Proportion
                            fixes_applied += 1
                
                elif len(row.cells) >= 3:  # Employee statement table
                    # Check if this is the employee statement table
                    if row_idx == 0 and 'STATEMENT' in str([cell.text for cell in row.cells]):
                        continue  # Skip header row
                    elif row_idx == 1 and 'Name of employee in full' in row.cells[1].text:
                        row.cells[2].text = employee_name  # Employee name
                        fixes_applied += 1
                    elif row_idx == 2 and 'Sex' in row.cells[1].text:
                        row.cells[2].text = gender  # Gender
                        fixes_applied += 1
                    elif row_idx == 3 and 'Religion' in row.cells[1].text:
                        row.cells[2].text = religion  # Religion
                        fixes_applied += 1
                    elif row_idx == 4 and 'Marital Status' in row.cells[1].text:
                        row.cells[2].text = marital_status  # Marital Status
                        fixes_applied += 1
                    elif row_idx == 5 and 'Department' in row.cells[1].text:
                        row.cells[2].text = department  # Department
                        fixes_applied += 1
                    elif row_idx == 6 and 'Post held' in row.cells[1].text:
                        row.cells[2].text = post_held  # Post held
                        fixes_applied += 1
                    elif row_idx == 7 and 'Date of appointment' in row.cells[1].text:
                        row.cells[2].text = appointment_date  # Date of appointment
                        fixes_applied += 1
                    elif row_idx == 8 and 'Permanent address' in row.cells[1].text:
                        row.cells[2].text = employee_address  # Address
                        fixes_applied += 1

        return fixes_applied

    def fill_epf_nomination_form(self, doc: Document, structure: Dict) -> int:
        """Fill EPF Nomination form with employee and nominee details."""
        print(f"  🔧 fill_epf_nomination_form called")
        print(f"  📷 Signature image path: {self.signature_image_path}")
        fixes_applied = 0
        personal_details = self.form_fields.get('personal_details', {})
        epf_gratuity = self.form_fields.get('epf_and_gratuity', {})
        
        # Employee details
        employee_name = personal_details.get('name', '')
        father_name = personal_details.get('father_name', '')
        dob = personal_details.get('date_of_birth', {})
        dob_value = dob.get('value', '') if isinstance(dob, dict) else str(dob)
        gender = personal_details.get('gender', '')
        employee_address = self.addr.get('current', {}).get('town_or_city_name', '') or self.addr.get('permanent', {}).get('town_or_city_name', '')
        
        # PF Account details
        pf_account_no = epf_gratuity.get('pf_account_no', '')
        marital_status = epf_gratuity.get('marital_status', '')
        
        # Nominee details
        nominee = epf_gratuity.get('nominee', {})
        nominee_name = nominee.get('name', '')
        nominee_relationship = nominee.get('relationship', '')
        nominee_address = nominee.get('address', '')
        nominee_dob = nominee.get('date_of_birth', {})
        nominee_dob_value = nominee_dob.get('value', '') if isinstance(nominee_dob, dict) else str(nominee_dob)
        nominee_share = nominee.get('share', '')
        
        # Family member details
        family_member = epf_gratuity.get('family_member_1', {})
        family_name = family_member.get('name', '')
        family_relationship = family_member.get('relationship', '')
        family_age = family_member.get('age', '')
        family_address = family_member.get('address', '')
        
        # Form details - use today's date
        from datetime import datetime
        today = datetime.now()
        sign_date = today.strftime('%d-%m-%Y')
        
        # Fill paragraphs with actual form patterns
        for paragraph in doc.paragraphs:
            text = paragraph.text
            # Name (IN BLOCK LETTERS) : \t
            if 'Name (IN BLOCK LETTERS)' in text:
                paragraph.text = f"Name (IN BLOCK LETTERS) : \t{employee_name}"
                fixes_applied += 1
            # Date of Birth :\t3. Account No. \t
            elif 'Date of Birth :' in text and 'Account No.' in text:
                paragraph.text = f"Date of Birth :\t{dob_value}\t3. Account No. \t{pf_account_no}"
                fixes_applied += 1
            # Marital Status
            elif 'Marital Status' in text and '4. *Sex : MALE/FEMALE:' in text:
                paragraph.text = f"4. *Sex : MALE/FEMALE:\t5. Marital Status \t{marital_status}"
                fixes_applied += 1
            # Address Permanent :\t
            elif 'Address Permanent' in text and text.strip().endswith('\t'):
                paragraph.text = f"6. Address Permanent :\t{employee_address}"
                fixes_applied += 1
            # Date \t (for signature date)
            elif text.strip() == 'Date \t':
                paragraph.text = f"Date \t{sign_date}"
                fixes_applied += 1
            # Employer certificate - Shri / Smt./ Miss
            elif 'Shri / Smt./ Miss' in text and 'employed in my establishment' in text:
                gender_prefix = 'Smt.' if gender.lower() in ['female', 'f'] else 'Shri'
                paragraph.text = text.replace('Shri / Smt./ Miss\t', f'{gender_prefix} {employee_name} ')
                fixes_applied += 1

        # Fill tables
        for table in doc.tables:
            for row_idx, row in enumerate(table.rows):
                if len(row.cells) >= 6:  # Nominee table
                    # Check if this is the nominee table (first row has headers)
                    if row_idx == 0 and 'Name of the Nominee' in str([cell.text for cell in row.cells]):
                        continue  # Skip header row
                    elif row_idx == 2 and not row.cells[0].text.strip():  # First empty nominee row
                        row.cells[0].text = nominee_name  # Nominee Name
                        row.cells[1].text = nominee_address  # Address
                        row.cells[2].text = nominee_relationship  # Relationship
                        row.cells[3].text = nominee_dob_value  # Date of Birth
                        row.cells[4].text = nominee_share  # Share
                        row.cells[5].text = ''  # Guardian (empty)
                        fixes_applied += 1
                
                elif len(row.cells) >= 4:  # Family member table
                    # Check if this is the family member table
                    if row_idx == 0 and 'Name & Address of the Family Member' in str([cell.text for cell in row.cells]):
                        continue  # Skip header row
                    elif row_idx == 1 and row.cells[0].text.strip() == '(1)':  # Header row
                        continue  # Skip header row
                    elif row_idx == 2 and not row.cells[0].text.strip():  # First family member row
                        row.cells[0].text = '1'  # Sr. No.
                        row.cells[1].text = f"{family_name}, {family_address}"  # Name & Address
                        row.cells[2].text = family_age  # Age
                        row.cells[3].text = family_relationship  # Relationship
                        fixes_applied += 1
                
                elif len(row.cells) >= 3:  # Third table (nominee details)
                    # Check if this is the third table with nominee details
                    if row_idx == 0 and 'Name and Address of the nominee' in str([cell.text for cell in row.cells]):
                        continue  # Skip header row
                    elif row_idx == 1 and not row.cells[0].text.strip():  # First nominee row
                        row.cells[0].text = f"{nominee_name}, {nominee_address}"  # Name and Address
                        row.cells[1].text = nominee_dob_value  # Date of Birth
                        row.cells[2].text = nominee_relationship  # Relationship with member
                        fixes_applied += 1

        # Add signature at the end of the document (EPF Nomination Form)
        # Look for signature fields in paragraphs - check ALL paragraphs, especially those with "employer" or "establishment"
        print(f"  🔍 Starting EPF Nomination signature insertion...")
        print(f"  📷 Signature image path available: {bool(self.signature_image_path)}")
        signature_inserted = False
        
        # First, look for the specific "Signature of the employer" pattern
        for i, p in enumerate(doc.paragraphs):
            text = (p.text or "").strip()
            text_lower = text.lower()
            
            # Check if paragraph already has an image
            has_image = False
            try:
                for run in p.runs:
                    if run._element.xpath('.//a:blip') or run._element.xpath('.//pic:pic'):
                        has_image = True
                        break
            except:
                pass
            
            if has_image:
                continue
            
            # Look for "Signature of the employer" or similar patterns
            if ("signature" in text_lower and 
                ("employer" in text_lower or "establishment" in text_lower or "authorised" in text_lower or "authorized" in text_lower)):
                # This paragraph contains the signature label
                print(f"  🔍 Found employer signature label: {text[:100]}")
                
                # Try multiple strategies:
                # 1. If paragraph is mostly just the label text (ends with the label), try inserting in same paragraph
                #    (some forms have label + empty space in same paragraph)
                if len(text) < 150:  # Reasonable length for a label paragraph
                    # Check if paragraph ends with signature label (might have space for signature after)
                    if text_lower.endswith("establishment") or text_lower.endswith("officer"):
                        # Try inserting in the same paragraph (will clear and add image)
                        if self._insert_signature_image(p):
                            fixes_applied += 1
                            signature_inserted = True
                            print(f"  ✅ Inserted signature in EPF Nomination Form (same paragraph as employer signature label)")
                            break
                
                # 2. Check next 5 paragraphs for empty signature field (increased from 3)
                for j in range(1, 6):
                    if i + j < len(doc.paragraphs):
                        next_p = doc.paragraphs[i + j]
                        next_text = (next_p.text or "").strip()
                        
                        # Check if next paragraph has image
                        next_has_image = False
                        try:
                            for run in next_p.runs:
                                if run._element.xpath('.//a:blip') or run._element.xpath('.//pic:pic'):
                                    next_has_image = True
                                    break
                        except:
                            pass
                        
                        if next_has_image:
                            continue
                        
                        # If next paragraph is empty or just has placeholder text, insert signature there
                        if not next_text or self._is_placeholder(next_text) or len(next_text) < 10:
                            if self._insert_signature_image(next_p):
                                fixes_applied += 1
                                signature_inserted = True
                                print(f"  ✅ Inserted signature in EPF Nomination Form (after employer signature label, paragraph {i+j+1})")
                                break
                
                if signature_inserted:
                    break
                
                # 3. If still not found, try inserting in the label paragraph itself (will replace text)
                if not signature_inserted:
                    print(f"  ⚠️  Trying to insert signature in label paragraph itself...")
                    if self._insert_signature_image(p):
                        fixes_applied += 1
                        signature_inserted = True
                        print(f"  ✅ Inserted signature in EPF Nomination Form (replaced employer signature label text)")
                        break
        
        # If not found yet, check last 15 paragraphs for any signature fields
        if not signature_inserted:
            for p in doc.paragraphs[-15:]:
                text = (p.text or "").strip().lower()
                # Check if paragraph already has an image
                has_image = False
                try:
                    for run in p.runs:
                        if run._element.xpath('.//a:blip') or run._element.xpath('.//pic:pic'):
                            has_image = True
                            break
                except:
                    pass
                
                if ("signature" in text and (":" in text or text == "signature")) and not has_image:
                    # Check if this paragraph is empty or just has "Signature:" text
                    if text in ["signature:", "signature"] or self._is_placeholder(text):
                        if self._insert_signature_image(p):
                            fixes_applied += 1
                            signature_inserted = True
                            print(f"  ✅ Inserted signature in EPF Nomination Form paragraph")
                            break
        
        # Also check tables for signature fields (especially at the end)
        # Look for "employer" or "establishment" signature in tables
        if not signature_inserted:
            for table in doc.tables:
                # Check ALL rows, not just last 3
                for row in table.rows:
                    for cell_idx, cell in enumerate(row.cells):
                        cell_text = (cell.text or "").strip().lower()
                        
                        # Look for employer/establishment signature in tables
                        if ("signature" in cell_text and 
                            ("employer" in cell_text or "establishment" in cell_text or "authorised" in cell_text or "authorized" in cell_text) and
                            not self._cell_has_image(cell)):
                            print(f"  🔍 Found employer signature in table cell: {cell_text[:50]}")
                            
                            # Try inserting in the same cell (will clear text)
                            if self._insert_signature_in_cell(cell):
                                fixes_applied += 1
                                signature_inserted = True
                                print(f"  ✅ Inserted signature in EPF Nomination Form table cell (employer signature)")
                                break
                            
                            # If same cell didn't work, try next cell
                            if not signature_inserted and cell_idx < len(row.cells) - 1:
                                next_cell = row.cells[cell_idx + 1]
                                if not self._cell_has_image(next_cell):
                                    if self._insert_signature_in_cell(next_cell):
                                        fixes_applied += 1
                                        signature_inserted = True
                                        print(f"  ✅ Inserted signature in EPF Nomination Form table (next cell after employer signature)")
                                        break
                    
                    if signature_inserted:
                        break
                
                if signature_inserted:
                    break

        return fixes_applied

    def fill_general_form(self, doc: Document, structure: Dict) -> int:
        fixes = 0
        name = self._get("name")
        email = self._get("email")
        addr_any = (self.addr.get("current") or self.addr.get("permanent") or {})
        address = addr_any.get("town_or_city_name", "")

        # paragraphs
        for p in doc.paragraphs:
            t = p.text
            if re.fullmatch(r"\s*Name:\s*", t):
                p.text = f"Name: {name}"; fixes += 1
            elif re.search(SmartFormPopulator.EMAIL_LABEL_RE, t or ""):
                label = t.split(":",1)[0] if ":" in (t or "") else "Email"
                p.text = f"{label.strip()}: {email}"; fixes += 1
            elif SmartFormPopulator.ADDRESS_LABEL_RE.search(t or ""):
                label = t.split(":",1)[0] if ":" in (t or "") else "Address"
                p.text = f"{label.strip()}: {address}"; fixes += 1

        # tables
        for table in doc.tables:
            for row in table.rows:
                if len(row.cells) >= 2:
                    L = (row.cells[0].text or "").strip()
                    # Check for signature fields
                    if "signature" in L.lower():
                        if self._insert_signature_in_cell(row.cells[-1]):
                            fixes += 1
                    elif re.fullmatch(r"\s*Name\s*:?\s*", L, re.I):
                        row.cells[-1].text = name; fixes += 1
                    elif SmartFormPopulator.EMAIL_LABEL_RE.search(L):
                        row.cells[-1].text = email; fixes += 1
                    elif SmartFormPopulator.ADDRESS_LABEL_RE.search(L):
                        row.cells[-1].text = address; fixes += 1
        return fixes

    # ----------------------------
    # Batch
    # ----------------------------
    def populate_all_forms(self, templates_dir: str, output_dir: str) -> int:
        print("🚀 Starting Smart Form Population")
        print(f"📁 Templates: {templates_dir}")
        print(f"📁 Output:    {output_dir}")

        os.makedirs(output_dir, exist_ok=True)
        template_files = [f for f in os.listdir(templates_dir) if f.lower().endswith(".docx")]
        if not template_files:
            print("❌ No template files found!")
            return 0

        print(f"📄 Found {len(template_files)} templates")
        ok = 0
        for name in template_files:
            src = os.path.join(templates_dir, name)
            dst = os.path.join(output_dir, f"smart_{name}")
            try:
                if self.populate_form_smart(src, dst):
                    ok += 1
            except Exception as e:
                import traceback
                print(f"❌ Error processing {name}: {e}")
                print(traceback.format_exc())
        print(f"\n🎉 Completed! {ok}/{len(template_files)} forms populated successfully")
        return ok


def main():
    import sys
    if len(sys.argv) < 3:
        print("Usage: python smart_form_populator.py <data_file> <templates_dir> [output_dir]")
        sys.exit(1)
    data_file = sys.argv[1]
    templates_dir = sys.argv[2]
    output_dir = sys.argv[3] if len(sys.argv) > 3 else "populated_forms_smart"

    pop = SmartFormPopulator(data_file)
    count = pop.populate_all_forms(templates_dir, output_dir)
    sys.exit(0 if count > 0 else 2)


if __name__ == "__main__":
    main()
